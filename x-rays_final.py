import logging
import secrets
import os
from flask import Flask, render_template, request, session, make_response, jsonify, send_file
from spellchecker import SpellChecker
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF
from Database.db_ps import create_messages_table, save_chat, get_chat_history, clear_all_chats, save_appointment
from Database.db_connection import close_db
from static.src.helper import download_hugging_face_embeddings
from langchain_pinecone import PineconeVectorStore
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from static.src.prompt import *
from Database.db_connection import *
from Database.db_ps import *
from langchain.embeddings import HuggingFaceEmbeddings
import requests  # Import the requests library
import re
import random
# PyTorch and Image Processing Imports
import torch
import torch.nn as nn
import torch.nn.functional as F
import torchvision.models as models
from torchvision import transforms
from PIL import Image
import io
import pickle

# FastAPI Imports (needed for image processing within Flask)
from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
# from fastapi.middleware import asgi_request_id
from transformers import AutoModelForSequenceClassification, AutoTokenizer, pipeline
from werkzeug.utils import secure_filename

# Move your FastAPI image classification code to a separate file (api.py):

# api.py
from fastapi import FastAPI, UploadFile, File
import torch
# ... [rest of your FastAPI image classification code]

# Keep Flask app in main file (app.py)
from flask import Flask
# ... [rest of your Flask code]
# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", secrets.token_hex(32))
# app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize the database
with app.app_context():
    create_messages_table()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Pinecone and LLM
PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Use Hugging Face 1024-d embeddings for Pinecone
embeddings = HuggingFaceEmbeddings(model_name="intfloat/e5-large")
docresearch = PineconeVectorStore.from_existing_index("curenow-chatbot", embedding=embeddings)

llm = ChatGoogleGenerativeAI(
    model="gemini-2.0-flash-lite-preview-02-05",
    temperature=0.7,
    google_api_key=GEMINI_API_KEY,
    max_retries=2,
    max_tokens=2000,
)

prompt = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
    ("human", "{input}"),
])

retriever = docresearch.as_retriever(search_type='similarity', search_kwargs={"k": 3})
# retriever = docresearch.as_retriever(search_type='mmr', search_kwargs={"k": 3, "lambda_mult": 0.7})

question_answer_chain = create_stuff_documents_chain(llm, prompt)
rag_chain = create_retrieval_chain(retriever, question_answer_chain)

# Spell Checker Instance
spell = SpellChecker()

#======================================================================================
#MODEL INTEGRATION
##======================================================================================

# Image-Based Disease Prediction Model
class MultiHeadMedicalModel(nn.Module):
    def __init__(self, num_brain_classes=4, num_lung_classes=4, num_skin_classes=9):
        super().__init__()
        # Backbone for grayscale images (brain & lung)
        self.backbone_gray = models.efficientnet_b0(weights=models.EfficientNet_B0_Weights.DEFAULT)
        self.backbone_gray.features[0][0] = nn.Conv2d(1, 32, kernel_size=3, stride=2, padding=1, bias=False)
        # Backbone for RGB images (skin)
        self.backbone_rgb = models.efficientnet_b0(weights=models.EfficientNet_B0_Weights.DEFAULT)

        # Freeze early layers, unfreeze last 3 layers
        for param in self.backbone_gray.parameters():
            param.requires_grad = False
        for param in self.backbone_rgb.parameters():
            param.requires_grad = False
        for param in self.backbone_gray.features[-3:].parameters():
            param.requires_grad = True
        for param in self.backbone_rgb.features[-3:].parameters():
            param.requires_grad = True

        # Classification heads
        self.brain_classifier = nn.Sequential(
            nn.Dropout(0.25),
            nn.Linear(1280, num_brain_classes)  # Use index 1 for weight keys to match the saved state_dict
        )

        self.lung_classifier = nn.Sequential(
            nn.Dropout(0.25),
            nn.Linear(1280, num_lung_classes)
        )

        self.skin_classifier = nn.Sequential(
            nn.Dropout(0.25),
            nn.Linear(1280, num_skin_classes)
        )

    def forward(self, x, task):
        if task in ["brain", "lung"]:
            x = self.backbone_gray.features(x)
            x = self.backbone_gray.avgpool(x)
        else:
            x = self.backbone_rgb.features(x)
            x = self.backbone_rgb.avgpool(x)

        x = torch.flatten(x, 1)
        return {
            "brain": self.brain_classifier(x),
            "lung": self.lung_classifier(x),
            "skin": self.skin_classifier(x)
        }[task]

# Load model
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model = MultiHeadMedicalModel(num_brain_classes=4, num_lung_classes=4, num_skin_classes=9).to(device)
model.load_state_dict(torch.load("x-rayy/Multi_head_medical_model.pth", map_location=device))  # Ensure correct path
model.eval()

# Labels
task_labels = {
    "brain": ["Glioma", "Meningioma", "No Tumor", "Pituitary"],
    "lung": ["Corona Virus Disease", "Normal", "Pneumonia", "Tuberculosis"],
    "skin": ["Actinic keratosis", "Dermatofibroma", "Squamous cell carcinoma", "Atopic Dermatitis",
             "Melanocytic nevus", "Tinea Ringworm Candidiasis", "Benign keratosis", "Melanoma", "Vascular lesion"]
}

# Image Transformations
transform = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
])

def predict_image(task: str, image_file): #image_bytes: bytes
    try:
        if task not in task_labels:
            return {"error": f"Invalid task. Choose from: {list(task_labels.keys())}"}

        image = Image.open(io.BytesIO(image_file)) #io.BytesIO(image_bytes)
        if task in ["brain", "lung"]:
            image = image.convert("L")  # Convert to grayscale

        image = transform(image).unsqueeze(0).to(device)

        with torch.no_grad():
            output = model(image, task)
            probs = F.softmax(output, dim=1)
            predicted_label = task_labels[task][torch.argmax(probs).item()]

        return {"task": task, "predicted_class": predicted_label}
    except Exception as e:
        logger.error(f"Error processing image: {e}")
        return {"error": f"Error processing the image: {e}"}

#======================================================================================

@app.route("/")
def index():
    session_id = request.cookies.get("session_id", secrets.token_hex(16))
    chat_history = get_chat_history(session_id) or []

    formatted_history = [
        {
            "user": msg["user_message"],
            "bot": msg["bot_response"],
            "timestamp": msg["timestamp"].strftime("%I:%M %p") if msg["timestamp"] else "Unknown"
        }
        for msg in chat_history
    ]

    response = make_response(render_template("x-rat_final.html", chat_history=formatted_history))
    response.set_cookie("session_id", session_id, max_age=60 * 60 * 24 * 7)  # Cookie valid for 7 days
    return response

import re

def clean_and_format_response(bot_response):
    """Auto-corrects and formats AI-generated text for clarity."""

    # Remove unwanted introduction phrases
    unwanted_phrases = [
        "The text discusses",
        "The document provides information on",
        "This passage explains",
        "This section talks about",
        "The provided text"
    ]
    for phrase in unwanted_phrases:
        bot_response = bot_response.replace(phrase, "")

    # Ensure section headers are properly formatted
    bot_response = re.sub(r"\b(Definition|Causes|Treatment|When to Seek Medical Attention|Symptoms)\b", r"<b>\1</b>", bot_response)

    # Convert Markdown-style bold (**text**) to HTML bold
    bot_response = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", bot_response)

    # Ensure list formatting
    bot_response = bot_response.replace("* ", "- ").replace("\n", "<br>")

    return bot_response

@app.route("/get", methods=["POST"])
def chat():
    session_id = request.cookies.get("session_id", secrets.token_hex(16))
    msg = request.form.get("msg", "").strip().lower()  # Convert input to lowercase

    if not msg:
        return jsonify({"error": "Message cannot be empty."}), 400

    logger.info(f"Received message: {msg}")
    # responses_of_bot = ["Hello!How can"]

    # Handle greetings separately
    greetings = ["hi", "hii","helloo","hello", "hey", "hola", "hi there", "good morning", "good evening"]
    greeting_responses = [
        "Hello! How can I assist you today?",
        "Hi there! What can I do for you?",
        "Hey! How's your day going?",
        "Hola! How can I help you?",
        "Good to see you! What do you need help with?"
    ]
    
    if any(greet in msg for greet in greetings):
        bot_response = random.choice(greeting_responses)
        save_chat(session_id, msg, bot_response)
        return jsonify({"user": msg, "bot": bot_response})

    # Auto-correct message spelling
    msg_corrected = ' '.join([spell.correction(word) or word for word in msg.split()])
    if not msg_corrected.strip():
        return jsonify({"error": "Invalid request"}), 400

    # Get response from retrieval model
    response = rag_chain.invoke({
        "input": msg_corrected,
        "format": "structured",
        "context": "Provide a well-formatted, user-friendly response with bullet points, markdown, and brief explanations."
    })

    # Ensure response structure
    if isinstance(response, dict):
        bot_response = response.get("answer", "Sorry, I couldn't process that request.")
    else:
        bot_response = "Sorry, I couldn't process that request."

    # Remove response if it contains "Not enough data available."
    if "Not enough data available" in bot_response:
        return jsonify({"user": msg, "bot": ""})  # Send an empty response

    bot_response = clean_and_format_response(bot_response)

    save_chat(session_id, msg, bot_response)
    return jsonify({"user": msg, "bot": bot_response})
import base64
from io import BytesIO


@app.route("/get_response", methods=["POST"])
def get_response():
    session_id = request.cookies.get("session_id", secrets.token_hex(16))
    msg = request.form.get("msg", "").strip().lower() if 'msg' in request.form else ""

    if 'xray_image' in request.files:
        image = request.files['xray_image']
        task = request.form.get('task')
        
        if not image:
            return jsonify({"error": "No image provided."}), 400
        if not task:
            return jsonify({"error": "No task provided."}), 400
        
        try:
            filename = secure_filename(image.filename)
            image_bytes = image.read()
            
            # Convert image to Base64
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            prediction = predict_image(task=task, image_file=image_bytes)

            if "predicted_class" in prediction:
                predicted_class = prediction["predicted_class"]
                bot_response = f"Based on the X-ray, the predicted condition is: {predicted_class}"
                save_chat(session_id, "X-ray uploaded", bot_response)
                
                return jsonify({
                    "user": "X-ray uploaded",
                    "image": image_base64,
                    "bot": bot_response
                      # Return the Base64-encoded image
                })
            else:
                return jsonify({"error": prediction.get("error", "Failed to get prediction.")}), 500

        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return jsonify({"error": f"Error processing the image: {e}"}), 500

    elif msg:
        msg_corrected = ' '.join([spell.correction(word) or word for word in msg.split()])
        if not msg_corrected.strip():
            return jsonify({"error": "Invalid request"}), 400

        response = rag_chain.invoke({
            "input": msg_corrected,
            "format": "structured",
            "context": "Provide a well-formatted, user-friendly response with bullet points, markdown, and brief explanations."
        })

        bot_response = response.get("answer", "Sorry, I couldn't process that request.") if isinstance(response, dict) else "Sorry, I couldn't process that request."

        if "Not enough data available" in bot_response:
            return jsonify({"user": msg, "bot": ""})

        bot_response = clean_and_format_response(bot_response)
        save_chat(session_id, msg, bot_response)
        return jsonify({"user": msg, "bot": bot_response})
    
    return jsonify({"error": "No message or image provided."}), 400



@app.route("/book_appointment", methods=["POST"])
def book_appointment():
    session_id = request.cookies.get("session_id", secrets.token_hex(16))

    # Collect the form data
    patient_name = request.form.get("patient_name")
    age = request.form.get("age")
    gender = request.form.get("gender")
    symptoms = request.form.get("symptoms")
    doctor_id = request.form.get("doctor_id")
    appointment_date = request.form.get("appointment_date")

    # Validate input fields
    if not patient_name or not age or not gender or not symptoms or not doctor_id or not appointment_date:
        return jsonify({"error": "Please provide all details to book the appointment."}), 400

    # Save the appointment to the database
    save_appointment(patient_name, age, gender, symptoms, doctor_id, appointment_date)

    bot_response = f"Your appointment with Doctor ID {doctor_id} has been successfully booked for {appointment_date}. See you soon!"
    save_chat(session_id, "Book appointment", bot_response)

    return jsonify({"user": "Book appointment", "bot": bot_response})

@app.route("/clear_chats", methods=["POST"])
def clear_chats():
    return jsonify({"success": clear_all_chats()}), 200

@app.route("/export_chat", methods=["GET"])
def export_chat():
    session_id = request.cookies.get("session_id", secrets.token_hex(16))
    format_type = request.args.get("format", "txt")
    chat_history = get_chat_history(session_id)
    if not chat_history:
        return "No chat history to export.", 400

    filename = f"chat_history.{format_type}"

    if format_type == "txt":
        with open(filename, "w", encoding="utf-8") as f:
            for chat in chat_history:
                f.write(f"User: {chat['user_message']}\n")
                f.write(f"Bot: {chat['bot_response']}\n\n")

    elif format_type == "docx":
        doc = Document()
        for chat in chat_history:
            p = doc.add_paragraph()
            p.add_run(f"User: {chat['user_message']}\n").bold = True
            p.add_run(f"Bot: {chat['bot_response']}\n\n")
        doc.save(filename)

    elif format_type == "pdf":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        font_path = os.path.abspath("fonts/DejaVuSans.ttf")
        if not os.path.exists(font_path):
            return "Font file not found.", 500

        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)

        for chat in chat_history:
            pdf.multi_cell(0, 10, f"User: {chat['user_message']}", align="L")
            pdf.multi_cell(0, 10, f"Bot: {chat['bot_response']}", align="L")
            pdf.cell(0, 10, "", ln=True)

        pdf.output(filename, "F")

    else:
        return "Invalid format.", 400

    return send_file(filename, as_attachment=True)


# @app.route("/upload", methods=["POST"])
# def upload_file():
#     session_id = request.cookies.get("session_id", secrets.token_hex(16))
#     msg = request.form.get("msg", "").strip().lower() if 'msg' in request.form else ""
    
#     if "file" not in request.files:
#         return jsonify({"error": "No file part"}), 400

#     file = request.files["file"]

#     if file.filename == "":
#         return jsonify({"error": "No selected file"}), 400

#     if file:
#         filename = secure_filename(file.filename)
#         filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        
#         # Save the uploaded file
#         file.save(filepath)
        
#         # Retrieve the latest uploaded file
#         latest_file = get_latest_file(app.config["UPLOAD_FOLDER"])
        
#         return jsonify({"message": "File uploaded successfully", "latest_file": latest_file}), 200
#     elif msg:
#         msg_corrected = ' '.join([spell.correction(word) or word for word in msg.split()])
#         if not msg_corrected.strip():
#             return jsonify({"error": "Invalid request"}), 400

#         response = rag_chain.invoke({
#             "input": msg_corrected,
#             "format": "structured",
#             "context": "Provide a well-formatted, user-friendly response with bullet points, markdown, and brief explanations."
#         })

#         bot_response = response.get("answer", "Sorry, I couldn't process that request.") if isinstance(response, dict) else "Sorry, I couldn't process that request."

#         if "Not enough data available" in bot_response:
#             return jsonify({"user": msg, "bot": ""})

#         bot_response = clean_and_format_response(bot_response)
#         save_chat(session_id, msg, bot_response)
#         return jsonify({"user": msg, "bot": bot_response})
    
#     return jsonify({"error": "No message or image provided."}), 400


# def get_latest_file(directory):
#     """Fetch the latest file from the uploads directory."""
#     files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
#     if not files:
#         return None
#     latest_file = max(files, key=lambda f: os.path.getctime(os.path.join(directory, f)))
#     return os.path.join(directory, latest_file)

# @app.route("/latest-image", methods=["GET"])
# def latest_image():
#     """Serve the latest uploaded image."""
#     latest_file = get_latest_file(app.config["UPLOAD_FOLDER"])
    
#     if latest_file:
#         return send_file(latest_file, mimetype="image/jpeg")
#     else:
#         return jsonify({"error": "No images found"}), 404

import os
import base64
import secrets
from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename




@app.route("/upload", methods=["POST"])
def upload_file():
    """Handles file uploads and provides predictions."""
    session_id = request.cookies.get("session_id", secrets.token_hex(16))
    msg = request.form.get("msg", "").strip().lower() if 'msg' in request.form else ""


    if "file" not in request.files:
        return jsonify({"error": "No file part in request."}), 400

    file = request.files["file"]
    
    if file.filename == "":
        return jsonify({"error": "No selected file."}), 400

    # Ensure task is provided
    task = request.form.get("task", "").strip()
    if not task:
        return jsonify({"error": "Task parameter missing."}), 400

    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)

        # Save the uploaded file
        file.save(filepath)

        # Read the saved image in binary mode
        with open(filepath, "rb") as img_file:
            image_bytes = img_file.read()

        # Convert image to Base64
        image_base64 = base64.b64encode(image_bytes).decode("utf-8")

        # Run prediction on the saved image
        prediction = predict_image(task=task, image_file=image_bytes)

        if "predicted_class" in prediction:
            predicted_class = prediction["predicted_class"]
            bot_response = f"Based on the X-ray, the predicted condition is: {predicted_class}"
            save_chat(session_id, "X-ray uploaded", bot_response)

            return jsonify({
                "user": "X-ray uploaded",
                "bot": bot_response,
                "image": image_base64,  # Return Base64-encoded image
                "filename": filename  # Filename for reference
            })
        else:return jsonify({"error": "Prediction failed.", "details": prediction}), 500
        
    elif msg:
        msg_corrected = ' '.join([spell.correction(word) or word for word in msg.split()])
        if not msg_corrected.strip():
            return jsonify({"error": "Invalid request"}), 400

        response = rag_chain.invoke({
            "input": msg_corrected,
            "format": "structured",
            "context": "Provide a well-formatted, user-friendly response with bullet points, markdown, and brief explanations."
        })

        bot_response = response.get("answer", "Sorry, I couldn't process that request.") if isinstance(response, dict) else "Sorry, I couldn't process that request."

        if "Not enough data available" in bot_response:
            return jsonify({"user": msg, "bot": ""})

        bot_response = clean_and_format_response(bot_response)
        save_chat(session_id, msg, bot_response)
        return jsonify({"user": msg, "bot": bot_response})

       

    return jsonify({"error": "File processing error"}), 400


@app.route("/latest-image", methods=["GET"])
def latest_image():
    """Serve the latest uploaded image."""
    latest_file = get_latest_file(app.config["UPLOAD_FOLDER"])
    
    if latest_file:
        return send_file(latest_file, mimetype="image/jpeg")
    else:
        return jsonify({"error": "No images found"}), 404


def get_latest_file(directory):
    """Fetch the latest file from the uploads directory."""
    files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
    if not files:
        return None
    latest_file = max(files, key=lambda f: os.path.getctime(os.path.join(directory, f)))
    return os.path.join(directory, latest_file)



@app.teardown_appcontext
def teardown_db(exception):
    close_db(exception)

if __name__ == "__main__":
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(host="0.0.0.0", port=8080, debug=True)
